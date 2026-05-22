# read and print the text in a microsoft word document docx

import os
# read and print the text in a microsoft word document docx
from docx import Document
import re
import enchant
from enchant.checker import SpellChecker
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

import nltk
from nltk.corpus import stopwords

from nltk.tokenize import word_tokenize
from nltk.tokenize import sent_tokenize
from nltk.tokenize import regexp_tokenize

working_directory = os.getcwd()

# read the document
def read_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

# print the document
def print_docx(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        print(para.text)
    print("\n")

# create empty xlsx document with tab name
def create_empty_xlsx(file_path, sheet_name):
    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)
    print(f"Created empty xlsx file: {file_path} with sheet name: {sheet_name}")

# close all files and exit program
def close_all_files():
    print("\n ... exiting program")
    exit()


# create an excel document with 3 columns and add 3 rows of bakery recipes
import pandas
def create_excel_file(file_path):
    # create a pandas dataframe
    df = pd.DataFrame({
        'Recipe': ['Bakery Recipe 1', 'Bakery Recipe 2', 'Bakery Recipe 3'],
        'Ingredients': ['Flour, Sugar, Eggs', 'Flour, Sugar, Milk', 'Flour, Sugar, Butter'],
        'Instructions': ['Mix ingredients and bake', 'Mix ingredients and bake', 'Mix ingredients and bake']
    })
    # write the dataframe to an excel file
    df.to_excel(file_path, index=False)
    print(f"Created excel file: {file_path}")



# read text from a pdf file
def read_pdf(file_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text

# print the text from a pdf file
def print_pdf(file_path):
    text = read_pdf(file_path)
    print(text)

# read text from a pdf file and append to xlsx columns
def append_pdf_to_xlsx(file_path, xlsx_file_path, sheet_name):
    text = read_pdf(file_path)
    df = pd.read_excel(xlsx_file_path, sheet_name=sheet_name)
    df = df.append({'Text': text}, ignore_index=True)
    df.to_excel(xlsx_file_path, sheet_name=sheet_name, index=False)
    print(f"Appended pdf text to xlsx file: {xlsx_file_path} in sheet: {sheet_name}")